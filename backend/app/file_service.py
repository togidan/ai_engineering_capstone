import logging
import io
from typing import Optional
from pathlib import Path

try:
    import PyPDF2
    import pdfplumber
    from docx import Document
    DEPENDENCIES_AVAILABLE = True
except ImportError:
    DEPENDENCIES_AVAILABLE = False
    logging.warning("File processing dependencies not installed. Only .txt files will be supported.")

logger = logging.getLogger(__name__)

class FileProcessingService:
    """Service for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_formats = ['.txt']
        if DEPENDENCIES_AVAILABLE:
            self.supported_formats.extend(['.pdf', '.docx', '.doc'])
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        suffix = Path(filename).suffix.lower()
        return suffix in self.supported_formats
    
    def extract_text(self, file_content: bytes, filename: str) -> Optional[str]:
        """
        Extract text from file content based on file extension
        Returns None if extraction fails
        """
        try:
            suffix = Path(filename).suffix.lower()
            
            if suffix == '.txt':
                return self._extract_text_from_txt(file_content)
            elif suffix == '.pdf' and DEPENDENCIES_AVAILABLE:
                return self._extract_text_from_pdf(file_content)
            elif suffix in ['.docx', '.doc'] and DEPENDENCIES_AVAILABLE:
                return self._extract_text_from_docx(file_content)
            else:
                logger.error(f"Unsupported file format: {suffix}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            return None
    
    def _extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, fallback to other encodings
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin1')
                except UnicodeDecodeError:
                    return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Failed to decode text file: {e}")
            raise
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file using pdfplumber (more reliable than PyPDF2)"""
        text_content = []
        
        try:
            # Use pdfplumber as primary method
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        
            if text_content:
                return '\n\n'.join(text_content)
                
            # Fallback to PyPDF2 if pdfplumber fails
            logger.warning("pdfplumber extraction failed, trying PyPDF2")
            return self._extract_text_from_pdf_pypdf2(file_content)
            
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
            return self._extract_text_from_pdf_pypdf2(file_content)
    
    def _extract_text_from_pdf_pypdf2(self, file_content: bytes) -> str:
        """Fallback PDF extraction using PyPDF2"""
        text_content = []
        
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)
                    
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
            raise
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

# Global file service instance
file_service = FileProcessingService()